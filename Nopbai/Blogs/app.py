from flask import Flask, render_template, Response, request, redirect
from db import edit_post1
from db import add_post, get_post
from datetime import datetime
from docx import Document
from docx.shared import Cm
app= Flask(__name__)

post_base= get_post()
post = {}
a=0
for p in post_base:
    a+=1
    post[str(a)] = p
doc = Document()
@app.route("/resignation-letter")
def regis():
    return render_template("resignation-letter.html")
@app.route("/resignation-letter", methods=["POST"])
def regiss():
    fullname = request.form["fullname"]
    reason = request.form["reason"]
    doc.add_heading('A letter send to you....', level=2)
    para= doc.add_paragraph("Hello world ! ")
    para.add_run(fullname).bold = True
    para.add_run(". End of the letter ! Byee... ")
    para.add_run(reason).italic = True
    doc.save('Blogs/cv1.docx')
    return redirect("/resignation-letter", code=302)


@app.route("/about")
def about():
    return render_template("about.html")

@app.route("/",methods=["GET"])
def index():
    return render_template("blog.html", post = post)

@app.route("/", methods=["POST"])
def new_movie():
    now = datetime.now()
    dt_string = now.strftime("%Y-%m-%d %H:%M:%S")
    time= "Post at: "+str(dt_string)
    title= request.form["title"]
    content= request.form["content"]
    b= str(len(post)+1)
    post[b]= {
        "time":time,
        "title":title,
        "content": content
    } 
    add_post(time,title,content)
    return redirect("/", code=302)

@app.route("/post/<post_id>", methods=["GET"])
def detail(post_id):
    post1 = post.get(post_id)
    title= post1["title"]
    a= post_id
    return Response(render_template("edit.html",post1=post1, title=title, a=a), status=404, mimetype="text/html")

@app.route("/post/<post_id>", methods=["POST"])
def edit_post(post_id):
    now = datetime.now()
    dt_string = now.strftime("%Y-%m-%d %H:%M:%S")
    time= "Post at: "+str(dt_string)
    title= request.form["title"]
    content= request.form["content"]
    b= str(post_id)
    post[str(post_id)]= {
        "time" : time,
        "title":title,
        "content": content
    } 
    edit_post1(b,time, title, content)
    return redirect("/", code=302)

if __name__== "__main__":
    app.run(debug=True)
